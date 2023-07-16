import requests
from bs4 import BeautifulSoup
import json
import csv
from docx import Document


def scrape_indeed_jobs(job_name, location, sort_by):
    base_url = 'https://www.indeed.com/jobs'
    params = {
        'q': job_name,
        'l': location,
        'sort': sort_by
    }

    response = requests.get(base_url, params=params)
    soup = BeautifulSoup(response.content, 'html.parser')

    job_listings = []

    for job in soup.select('.jobsearch-SerpJobCard'):
        title = job.select_one('.jobtitle').text.strip()
        company = job.select_one('.company').text.strip()
        location = job.select_one('.location').text.strip()
        summary = job.select_one('.summary').text.strip()
        salary = job.select_one('.salaryText').text.strip() if job.select_one('.salaryText') else ''

        job_listings.append({
            'Title': title,
            'Company': company,
            'Location': location,
            'Summary': summary,
            'Salary': salary
        })

    return job_listings


def print_job_listings(job_listings):
    for i, job in enumerate(job_listings, start=1):
        print(f'Job {i}:')
        print(f'Title: {job["Title"]}')
        print(f'Company: {job["Company"]}')
        print(f'Location: {job["Location"]}')
        print(f'Summary: {job["Summary"]}')
        print(f'Salary: {job["Salary"]}')
        print()


def save_job_listings(job_listings, filename, file_format):
    if file_format == 'json':
        with open(filename, 'w') as f:
            json.dump(job_listings, f, indent=4)
    elif file_format == 'csv':
        with open(filename, 'w', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=job_listings[0].keys())
            writer.writeheader()
            writer.writerows(job_listings)
    elif file_format == 'docx':
        doc = Document()
        for job in job_listings:
            doc.add_paragraph(f'Title: {job["Title"]}')
            doc.add_paragraph(f'Company: {job["Company"]}')
            doc.add_paragraph(f'Location: {job["Location"]}')
            doc.add_paragraph(f'Summary: {job["Summary"]}')
            doc.add_paragraph(f'Salary: {job["Salary"]}')
            doc.add_paragraph()
        doc.save(filename)
    elif file_format == 'text':
        with open(filename, 'w') as f:
            for job in job_listings:
                f.write(f'Title: {job["Title"]}\n')
                f.write(f'Company: {job["Company"]}\n')
                f.write(f'Location: {job["Location"]}\n')
                f.write(f'Summary: {job["Summary"]}\n')
                f.write(f'Salary: {job["Salary"]}\n')
                f.write('\n')


# Example usage:
job_name = input('Enter job name: ')
location = input('Enter location: ')
sort_by = input('Sort by (date/rank/salary): ')

job_listings = scrape_indeed_jobs(job_name, location, sort_by)

print_job_listings(job_listings)

# Save as JSON
save_job_listings(job_listings, 'job_listings.json', 'json')

# Save as CSV
save_job_listings(job_listings, 'job_listings.csv', 'csv')

# Save as DOCX
save_job_listings(job_listings, 'job_listings.docx', 'docx')

# Save as text
save_job_listings(job_listings, 'job_listings.txt', 'text')
      
